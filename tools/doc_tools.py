"""
Document Tools for DOCX Processing
- Format-preserving text replacement (run-level)
- Headers/footers support
- File validation
"""

import os
from typing import Optional, Tuple, List
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from loguru import logger

from config import MAX_FILE_SIZE_BYTES, SUPPORTED_EXTENSIONS


# ============================================
# VALIDATION
# ============================================


def validate_docx(file_path: str) -> Tuple[bool, str]:
    """
    Validate that file is a proper .docx and within size limits.

    Returns:
        (is_valid, error_message)
    """
    # Check file exists
    if not os.path.exists(file_path):
        return False, "File not found."

    # Check extension
    _, ext = os.path.splitext(file_path)
    if ext.lower() not in SUPPORTED_EXTENSIONS:
        return False, f"Invalid file type. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"

    # Check file size
    file_size = os.path.getsize(file_path)
    if file_size > MAX_FILE_SIZE_BYTES:
        max_mb = MAX_FILE_SIZE_BYTES / (1024 * 1024)
        return False, f"File too large. Maximum size: {max_mb}MB"

    # Check if it's a valid DOCX (can be opened)
    try:
        doc = Document(file_path)
        # Quick sanity check - try to access paragraphs
        _ = len(doc.paragraphs)
    except PackageNotFoundError:
        return False, "Invalid or corrupted DOCX file."
    except Exception as e:
        return False, f"Cannot read file: {str(e)}"

    return True, ""


# ============================================
# HELPER FUNCTIONS
# ============================================


def _get_all_paragraphs(doc: Document) -> List:
    """
    Get all paragraphs from document body, headers, and footers.
    """
    paragraphs = []

    # Body paragraphs
    paragraphs.extend(doc.paragraphs)

    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    # Header and footer paragraphs (all sections)
    for section in doc.sections:
        # Headers
        if section.header:
            paragraphs.extend(section.header.paragraphs)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)

        # Footers
        if section.footer:
            paragraphs.extend(section.footer.paragraphs)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)

        # First page header/footer (if different)
        if section.first_page_header:
            paragraphs.extend(section.first_page_header.paragraphs)
        if section.first_page_footer:
            paragraphs.extend(section.first_page_footer.paragraphs)

        # Even page header/footer (if different)
        if section.even_page_header:
            paragraphs.extend(section.even_page_header.paragraphs)
        if section.even_page_footer:
            paragraphs.extend(section.even_page_footer.paragraphs)

    return paragraphs


def _replace_in_paragraph_runs(paragraph, search_text: str, replace_text: str) -> int:
    """
    Replace text at the run level to preserve formatting.

    This handles cases where searched text might span multiple runs.
    Returns the number of replacements made in this paragraph.
    """
    # Get full paragraph text
    full_text = paragraph.text

    if search_text not in full_text:
        return 0

    # Count occurrences for return value
    count = full_text.count(search_text)

    # Strategy: Rebuild runs while preserving formatting
    # This is complex because text can span multiple runs

    # First, try simple case: search text is entirely within single runs
    simple_replaced = False
    for run in paragraph.runs:
        if search_text in run.text:
            run.text = run.text.replace(search_text, replace_text)
            simple_replaced = True

    if simple_replaced:
        return count

    # Complex case: text spans multiple runs
    # We need to find where the text starts and ends across runs

    # Build a map of character positions to runs
    runs = paragraph.runs
    if not runs:
        return 0

    # Concatenate all run texts and track positions
    combined_text = ""
    run_boundaries = []  # [(start_pos, end_pos, run_index), ...]

    for i, run in enumerate(runs):
        start = len(combined_text)
        combined_text += run.text
        end = len(combined_text)
        run_boundaries.append((start, end, i))

    # Find all occurrences of search_text
    new_combined = combined_text.replace(search_text, replace_text)

    if new_combined == combined_text:
        return 0

    # Clear all runs and redistribute the new text
    # Preserve formatting of first run for the replacement
    if runs:
        # Store formatting info from runs
        run_formats = []
        for run in runs:
            run_formats.append(
                {
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size,
                }
            )

        # Clear all runs except first
        first_run = runs[0]
        for run in runs[1:]:
            run.text = ""

        # Put all text in first run (preserves its formatting)
        first_run.text = new_combined

    return count


# ============================================
# MAIN FUNCTIONS
# ============================================


def replace_text_in_docx(
    file_path: str, search_text: str, replace_text: str
) -> Optional[str]:
    """
    Replace text in DOCX while preserving formatting.
    Searches in body, tables, headers, and footers.

    Args:
        file_path: Path to the DOCX file
        search_text: Text to find
        replace_text: Text to replace with

    Returns:
        Path to the new file if replacements were made, None otherwise
    """
    try:
        doc = Document(file_path)
        total_replacements = 0

        # Get all paragraphs (body + tables + headers + footers)
        all_paragraphs = _get_all_paragraphs(doc)

        for paragraph in all_paragraphs:
            replacements = _replace_in_paragraph_runs(
                paragraph, search_text, replace_text
            )
            total_replacements += replacements

        if total_replacements == 0:
            return None

        # Generate output filename
        base, ext = os.path.splitext(file_path)
        new_filename = f"{base}_revisi{ext}"

        doc.save(new_filename)
        logger.info(
            f"Replaced {total_replacements} occurrences, saved to {new_filename}"
        )

        return new_filename

    except Exception as e:
        logger.error(f"Error replacing text: {e}")
        return None


def get_occurrences_with_context(file_path: str, search_text: str) -> List[dict]:
    """
    Find all occurrences of text with the full sentence as context.

    Args:
        file_path: Path to the DOCX file
        search_text: Text to find

    Returns:
        List of dicts with occurrence info:
        [
            {
                "index": 0,
                "sentence": "This is the full sentence containing the search text.",
                "paragraph_index": 2
            },
            ...
        ]
    """
    import re

    try:
        doc = Document(file_path)
        occurrences = []
        occurrence_index = 0

        # Get all paragraphs
        all_paragraphs = _get_all_paragraphs(doc)

        for para_idx, paragraph in enumerate(all_paragraphs):
            para_text = paragraph.text
            if search_text not in para_text:
                continue

            # Split paragraph into sentences
            # Use regex to split on sentence-ending punctuation
            sentences = re.split(r"(?<=[.!?])\s+", para_text)

            for sentence in sentences:
                if search_text in sentence:
                    # Count how many times search_text appears in this sentence
                    count_in_sentence = sentence.count(search_text)
                    for _ in range(count_in_sentence):
                        occurrences.append(
                            {
                                "index": occurrence_index,
                                "sentence": sentence.strip(),
                                "paragraph_index": para_idx,
                            }
                        )
                        occurrence_index += 1

        return occurrences

    except Exception as e:
        logger.error(f"Error finding occurrences: {e}")
        return []


def count_text_in_docx(file_path: str, search_text: str) -> int:
    """
    Count occurrences of text in DOCX.
    Searches in body, tables, headers, and footers.

    Args:
        file_path: Path to the DOCX file
        search_text: Text to count

    Returns:
        Number of occurrences found
    """
    try:
        doc = Document(file_path)
        count = 0

        # Get all paragraphs (body + tables + headers + footers)
        all_paragraphs = _get_all_paragraphs(doc)

        for paragraph in all_paragraphs:
            if search_text in paragraph.text:
                count += paragraph.text.count(search_text)

        return count

    except Exception as e:
        logger.error(f"Error counting text: {e}")
        return 0


def read_docx_full_text(file_path: str) -> str:
    """
    Extract all text from DOCX document.
    Includes body, tables, headers, and footers.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Full text content of the document
    """
    try:
        doc = Document(file_path)
        full_text = []

        # Body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                ]
                if row_text:
                    full_text.append(" | ".join(row_text))

        # Headers and footers (all sections)
        for section in doc.sections:
            # Main header
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        full_text.append(f"[HEADER] {para.text}")

            # Main footer
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        full_text.append(f"[FOOTER] {para.text}")

        return "\n".join(full_text)

    except Exception as e:
        logger.error(f"Error reading document: {e}")
        return ""


def apply_multiple_fixes(
    file_path: str, fixes: List[dict]
) -> Tuple[Optional[str], int, int, List[dict], List[dict]]:
    """
    Apply multiple find/replace fixes to a document.

    Args:
        file_path: Path to the DOCX file
        fixes: List of {"search": str, "replace": str} dicts

    Returns:
        (new_file_path, applied_count, skipped_count, applied_list, skipped_list)
    """
    try:
        doc = Document(file_path)
        applied_count = 0
        skipped_count = 0
        applied_list = []
        skipped_list = []

        all_paragraphs = _get_all_paragraphs(doc)

        for fix in fixes:
            search = fix.get("search", "")
            replace = fix.get("replace", "")

            if not search:
                skipped_count += 1
                skipped_list.append(fix)
                continue

            fix_applied = False
            for paragraph in all_paragraphs:
                if _replace_in_paragraph_runs(paragraph, search, replace) > 0:
                    fix_applied = True

            if fix_applied:
                applied_count += 1
                applied_list.append(fix)
            else:
                skipped_count += 1
                skipped_list.append(fix)

        if applied_count == 0:
            return None, 0, len(fixes), [], fixes

        # Generate output filename
        base, ext = os.path.splitext(file_path)
        new_filename = f"{base}_revisi{ext}"

        doc.save(new_filename)
        logger.info(
            f"Applied {applied_count} fixes, skipped {skipped_count}, saved to {new_filename}"
        )

        return new_filename, applied_count, skipped_count, applied_list, skipped_list

    except Exception as e:
        logger.error(f"Error applying fixes: {e}")
        return None, 0, len(fixes), [], fixes
