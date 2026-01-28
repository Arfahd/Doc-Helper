"""
Pytest fixtures for Enterprise Doc Bot tests.
"""

import os
import sys
import pytest
from docx import Document

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


@pytest.fixture
def temp_docx(tmp_path):
    """Create a temporary DOCX file for testing."""
    doc_path = tmp_path / "test_document.docx"
    doc = Document()

    # Add content with intentional errors
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test document with some erors and typos.")
    doc.add_paragraph("Teh quick brown fox jumps over teh lazy dog.")
    doc.add_paragraph("We need to recieve this message.")

    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "teh value"
    table.rows[1].cells[1].text = "another value"

    doc.save(doc_path)
    return str(doc_path)


@pytest.fixture
def clean_docx(tmp_path):
    """Create a clean DOCX file without errors."""
    doc_path = tmp_path / "clean_document.docx"
    doc = Document()

    doc.add_heading("Clean Document", level=1)
    doc.add_paragraph("This document has no spelling errors.")
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")

    doc.save(doc_path)
    return str(doc_path)


@pytest.fixture
def empty_docx(tmp_path):
    """Create an empty DOCX file."""
    doc_path = tmp_path / "empty_document.docx"
    doc = Document()
    doc.save(doc_path)
    return str(doc_path)


@pytest.fixture
def sample_fixes():
    """Sample fixes for testing."""
    return [
        {"search": "erors", "replace": "errors"},
        {"search": "Teh", "replace": "The"},
        {"search": "teh", "replace": "the"},
        {"search": "recieve", "replace": "receive"},
    ]


@pytest.fixture
def sample_ai_response_with_fixes():
    """Sample AI response containing JSON fixes."""
    return """Here is my analysis of the document.

I found several issues:

1. Spelling error: "erors" should be "errors"
2. Spelling error: "teh" should be "the"

```json
[
    {"search": "erors", "replace": "errors"},
    {"search": "Teh", "replace": "The"},
    {"search": "teh", "replace": "the"}
]
```
"""


@pytest.fixture
def sample_ai_response_no_fixes():
    """Sample AI response with no fixes."""
    return """Here is my analysis of the document.

The document looks good! No issues found.

```json
[]
```
"""
